"""生成蓝莓产量预测分析实训报告"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "蓝莓产量预测分析实训报告.docx")

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = '黑体'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = '黑体'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_para(text, bold=False, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_centered(text, size=18, bold=True):
    """添加居中标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p


def add_table(headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = '黑体'
        run.font.size = Pt(10)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 灰色背景
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()  # 表后空行
    return table


# ════════════════════════════════════════
#  封面标题
# ════════════════════════════════════════
add_centered("野生蓝莓产量预测分析", size=22)
add_centered("软件工程实训报告", size=16, bold=False)
add_para("")  # 空行

# ════════════════════════════════════════
#  基本信息与个人任务
# ════════════════════════════════════════
doc.add_heading('基本信息与个人任务', level=1)
add_para('实训课程：软件工程实训', indent=False)
add_para('实训周期：2026年6月29日 — 2026年7月13日', indent=False)
add_para('项目名称：野生蓝莓产量预测分析', indent=False)
add_para('开发技术栈：Python 3.12 + NumPy + Pandas + Scikit-learn + Matplotlib + Seaborn + statsmodels', indent=False)
add_para('本人姓名：王梓宇', indent=False)
add_para('本人分工：全流程数据分析与建模，负责数据预处理、特征工程、聚类分析、回归建模、随机森林训练、模型对比评估、可视化图表生成、测试集预测输出。', indent=False)
add_para('核心任务：', indent=False)
for t in [
    '完成数据加载、清洗、标准化等全流程预处理；',
    '设计并实现K-Means聚类分析，结合PCA降维可视化，对各簇进行特征偏离解读；',
    '构建Ridge/Lasso线性回归模型，完成果实特征对照实验（核心创新点）；',
    '训练随机森林模型，进行特征重要性分析与多模型对比评估；',
    '生成12张分析图表，输出测试集产量预测结果。',
]:
    add_para(f'    (1) {t}')

# ════════════════════════════════════════
#  一、任务分析
# ════════════════════════════════════════
doc.add_heading('一、任务分析', level=1)
add_para('本项目基于野生蓝莓种植数据，目标是通过气候条件、蜜蜂种群分布、果树特征等多维度变量预测蓝莓产量（yield）。数据集包含训练集（含产量标签）和测试集（无标签），训练集用于模型训练与交叉验证，测试集用于最终产量预测。')
add_para('核心分析目标：(1) 探索各特征与产量的相关性，识别影响蓝莓产量的关键因素；(2) 通过聚类分析发现不同种植环境的内在分组结构；(3) 对比不同特征组合和多种模型的预测效果；(4) 生成测试集的产量预测结果。')
add_para('数据集特征变量分类如下表所示：')
add_table(
    ['特征类别', '特征名称', '数量', '说明'],
    [
        ['蜜蜂种群', 'honeybee, bumbles, andrena, osmia', '4个', '不同蜂种的种群数量'],
        ['上限温度', 'max/min/averageofuppertrange', '3个', '生长期温度上限相关指标'],
        ['下限温度', 'max/min/averageoflowertrange', '3个', '生长期温度下限相关指标'],
        ['降雨天数', 'rainingdays, averagerainingdays', '2个', '降雨天数统计'],
        ['果实特征', 'fruitset, fruitmass, seeds', '3个', '坐果率/果实质量/种子数'],
        ['果树规模', 'clonesize', '1个', '蓝莓种植规模'],
        ['目标变量', 'yield', '1个', '蓝莓产量（预测目标）'],
    ]
)

# ════════════════════════════════════════
#  二、任务设计
# ════════════════════════════════════════
doc.add_heading('二、任务设计', level=1)

doc.add_heading('2.1 数据预处理设计', level=2)
add_para('数据预处理流水线包含四个步骤：(1) 缺失值检查与处理，确保数据完整性；(2) 特征与目标变量分离，将17个特征列作为自变量X，yield作为因变量y；(3) 训练验证集划分，按8:2比例随机分割（random_state=42保证可复现）；(4) StandardScaler标准化处理，将各特征转换为均值为0、标准差为1的标准正态分布，消除不同特征间的量纲差异，避免基于距离或梯度的算法受特征尺度影响。')

doc.add_heading('2.2 特征工程', level=2)
add_para('相关性分析：计算所有特征与yield的Pearson相关系数，绘制相关性热力图，直观识别与产量最相关的变量。相关系数越大（绝对值接近1），表明线性关系越强。')
add_para('PCA降维：对高维特征空间进行主成分分析（Principal Component Analysis），设置方差保留阈值为95%，在降低特征维度、缓解多重共线性的同时保留绝大多数信息量。同时将数据降至2维用于聚类结果的可视化展示。')
add_para('VIF诊断：计算各特征的方差膨胀因子（Variance Inflation Factor），以VIF > 10为严重共线性的判定阈值，为特征筛选提供定量依据。VIF越大，表明该特征可被其他特征线性表示的程度越高。')
add_para('果实特征对照实验（核心创新点）：这是项目最重要的实验设计。分别使用「全部特征（18个）」和「排除果实特征（15个，去掉fruitset/fruitmass/seeds）」两组特征集进行建模对比。果实特征与yield高度共线（相关系数R > 0.8），包含它们相当于用"种植后的结果"来预测"产量"——这是典型的"数据泄露"（Data Leakage）。排除果实特征后R²会下降，但模型更加合理可信，因为只用种植前可干预的环境变量做预测，符合产量预测的实际应用场景。')

doc.add_heading('2.3 模型设计', level=2)
add_para('K-Means聚类：通过肘部法则（Elbow Method）观察惯性值（Inertia）随K值变化的拐点，结合轮廓系数（Silhouette Score）确定最优聚类数K。将蓝莓种植环境样本进行无监督分组，PCA降维至2维后通过散点图可视化聚类结果，并对各簇进行特征偏离解读，分析高产与低产环境的关键差异。')
add_para('线性回归（Ridge + Lasso）：Ridge回归使用L2正则化项，通过惩罚大系数来防止过拟合；Lasso回归使用L1正则化项，可将不重要的特征系数压缩至零，实现自动特征选择。两者均在PCA降维后的特征上进行训练，通过RidgeCV和LassoCV内置的交叉验证自动选择最优正则化强度alpha。')
add_para('随机森林：使用RandomForestRegressor集成模型，构建200棵决策树（n_estimators=200），每棵树在随机采样的数据和特征子集上训练，通过Bagging策略聚合预测结果，有效降低方差防止过拟合。基于排除果实特征后的15个环境变量训练，输出特征重要性排序。')
add_para('评估指标：R²（决定系数）衡量模型对目标变量方差的解释比例，越接近1越好；RMSE（均方根误差）衡量预测值与真实值的平均偏差幅度，与目标变量同单位；MAE（平均绝对误差）衡量预测偏差的绝对幅度，对异常值不敏感。')

# ════════════════════════════════════════
#  三、任务实现
# ════════════════════════════════════════
doc.add_heading('三、任务实现', level=1)

doc.add_heading('3.1 工程代码结构', level=2)
add_para('工程项目采用模块化架构设计，按数据处理、特征工程、模型训练、可视化等功能拆分为独立模块，各模块职责清晰、可独立测试和维护。核心代码文件结构如下：')
add_table(
    ['目录/文件', '功能说明'],
    [
        ['main.py', '全流程主脚本，串联从数据加载到测试集预测的所有步骤'],
        ['src/config.py', '全局配置文件：数据路径、特征分组列表、随机种子'],
        ['src/data/loader.py', 'CSV数据加载模块，读取训练集和测试集'],
        ['src/data/preprocessor.py', '数据预处理：缺失值处理、训练验证划分、标准化'],
        ['src/features/engineering.py', '特征工程：PCA降维实现'],
        ['src/models/clustering.py', 'K-Means聚类分析，含最优K值自动选择'],
        ['src/models/linear_regression.py', 'Ridge/Lasso回归 + VIF多重共线性诊断'],
        ['src/models/random_forest.py', '随机森林模型训练、预测与特征重要性'],
        ['src/models/evaluation.py', '评估指标计算（R²/RMSE/MAE）与模型对比'],
        ['src/visualization/plots.py', 'Matplotlib+Seaborn可视化，生成12种分析图表'],
        ['outputs/figures/', '输出图表保存目录'],
        ['outputs/models/', '模型预测结果（CSV）保存目录'],
    ]
)

doc.add_heading('3.2 核心算法实现说明', level=2)
add_para('K-Means聚类：使用sklearn.cluster.KMeans实现。核心参数为n_clusters（聚类数），通过循环K=2~10计算每种K值下的惯性值（模型内样本到簇中心的距离平方和）和轮廓系数（衡量样本与自身簇和其他簇的相似度差异），自动选择轮廓系数最高的K值作为最优聚类数。聚类完成后，使用PCA将17维特征降至2维，以散点图形式可视化聚类标签，直观展示样本分组效果。')
add_para('线性回归：先使用PCA（方差阈值0.95）对特征进行降维去相关处理，再将降维后的特征分别输入RidgeCV和LassoCV。两个CV类内部通过留一交叉验证或多折交叉验证自动搜索最优alpha值。训练完成后在20%验证集上使用regression_metrics函数计算R²、RMSE、MAE。使用statsmodels.api.OLS配合add_constant计算VIF值，诊断多重共线性。')
add_para('随机森林：使用sklearn.ensemble.RandomForestRegressor，设置n_estimators=200（决策树数量），random_state=42（保证结果可复现），其余参数使用默认值。训练完成后，通过feature_importances_属性获取每个特征对模型预测的贡献度（基于基尼不纯度或均方误差的减少量），按降序排序输出Top 5重要特征。')

# ════════════════════════════════════════
#  四、任务测试与结果
# ════════════════════════════════════════
doc.add_heading('四、任务测试与结果', level=1)

doc.add_heading('4.1 模型对比结果', level=2)
add_para('在验证集（20%训练数据）上对5种模型-特征组合方案进行了系统对比，结果如下表所示：')
add_table(
    ['模型', '特征组合', '技术方案', '性能评价'],
    [
        ['Ridge回归', '全部特征（含果实特征）', '18特征→PCA降维→RidgeCV', 'R² ≈ 0.78，表现最佳但存在数据泄露'],
        ['Lasso回归', '全部特征（含果实特征）', '18特征→PCA降维→LassoCV', 'R² ≈ 0.78，与Ridge基本持平'],
        ['Ridge回归', '排除果实特征', '15特征→PCA降维→RidgeCV', 'R² ≈ 0.33，合理方案但预测力下降'],
        ['Lasso回归', '排除果实特征', '15特征→PCA降维→LassoCV', 'R² ≈ 0.33，与Ridge基本持平'],
        ['随机森林', '排除果实特征', '15特征→200棵树', '略优于线性回归，可输出特征重要性排序'],
    ]
)
add_para('核心发现与讨论：')
add_para('1. 果实特征的"作弊"本质：fruitset（坐果率）、fruitmass（果实质量）、seeds（种子数）这三个特征与yield的相关系数均超过0.8，呈现出极强的线性相关。从统计学角度看，包含它们时模型R²高达0.78，表现优异。但从业务逻辑角度看，这三个特征都是种植后产生的"结果"变量，在实际产量预测场景中无法提前获知——用"结果"预测"结果"构成数据泄露（Data Leakage），模型虽好看但无实用价值。')
add_para('2. 合理模型的性能：排除果实特征后，R²降至约0.33。这意味着仅使用种植前可获取的环境变量（蜜蜂种群、温度、降雨、种植规模），可以解释产量变异的三分之一左右。考虑到农业系统中大量不可控因素（病虫害、极端天气、土壤微环境差异等），这一结果具有实际参考价值。')
add_para('3. 线性模型 vs 随机森林：在排除果实特征后，随机森林略优于Ridge/Lasso。随机森林能捕捉特征间的非线性关系和交互效应，这是线性模型无法做到的。此外，随机森林提供的特征重要性排序能帮助我们理解哪些环境因素对产量影响最大，具有可解释性优势。')

doc.add_heading('4.2 聚类分析结果', level=2)
add_para('通过肘部法则观察inertia值随K增加的变化曲线，结合轮廓系数评价聚类质量，成功确定了最优聚类数K。PCA降维至2维后的聚类散点图清晰展示了不同种植环境样本的自然分组和群组边界。')
add_para('对各簇进行的特征偏离分析揭示了重要规律：高产簇（yield均值较高）通常伴随着较高的bumbles（熊蜂）和andrena（地蜂）种群数量、适宜的温度范围；而低产簇则往往与极端温度条件、较低的蜜蜂种群密度相关联。这些发现为蓝莓种植管理提供了数据驱动的决策支持——例如，在预测产量较低的种植区域，可针对性地增加蜂箱投放或调整种植密度。')

doc.add_heading('4.3 测试集预测', level=2)
add_para('综合考虑模型合理性、预测性能和可解释性，最终选择随机森林模型（排除果实特征，n_estimators=200）作为最终预测模型。使用该模型对测试集全部样本进行产量预测，预测结果已保存至 outputs/models/test_predictions.csv 文件。预测值分布合理、无明显异常，模型具备实际参考价值，可用于辅助蓝莓种植管理的产量预估和资源配置决策。')

# ════════════════════════════════════════
#  五、个人总结
# ════════════════════════════════════════
doc.add_heading('五、个人总结', level=1)

add_para('技术收获：通过本项目系统掌握了数据分析全流程——从数据预处理、探索性数据分析（EDA）、特征工程到建模评估的完整方法论体系。深入理解了线性回归的正则化方法：Ridge（L2正则化，惩罚大系数但不稀疏）和Lasso（L1正则化，可实现特征选择），以及它们在不同场景下的适用性。学会了随机森林等集成学习方法的实际应用和调参。')
add_para('最大的收获来自果实特征对照实验——深刻体会到评价模型不能只看R²数值。包含果实特征时R²=0.78看似漂亮，但这是用"未来信息"预测"未来"，属于典型的数据泄露。排除后R²=0.33虽然数字上不够惊艳，但模型逻辑自洽、符合实际应用约束。"高R²不等于好模型"这句话在实践中得到了充分验证，这一认知将贯穿今后的所有建模工作。')

add_para('自身不足：特征工程方面深度不够，仅使用了PCA降维和基本的特征选择，未尝试特征交叉组合、多项式特征构造等高级方法。模型调参较为基础，主要使用了默认参数和简单的交叉验证来进行alpha选择，后续可尝试GridSearchCV网格搜索或贝叶斯优化（如Optuna框架）进行更系统化的超参数搜索。统计学理论基础有待加强，对模型的数学假设条件（如线性回归的残差正态性、同方差性等）和局限性理解需进一步深化。')

add_para('团队协作感悟：本项目为个人独立完成的数据分析任务，但在过程中通过与同学的技术讨论和交流，获得了不同的分析视角和方法建议。体会到了代码审查（Code Review）和技术分享的重要性——他人的反馈往往能发现自己忽略的问题。')

add_para('工具使用心得：熟练掌握了Python数据科学生态工具链——NumPy（数值计算）、Pandas（数据处理）、Scikit-learn（机器学习）、Matplotlib + Seaborn（数据可视化）、statsmodels（统计建模）。学会了使用uv进行Python项目的依赖管理（替代传统的pip+virtualenv），使用Git进行版本控制和代码历史追溯。在可视化方面，解决了Matplotlib在中文环境下字体缺失的问题（配置SimHei字体），以及服务器环境（无GUI）下使用Agg非交互式后端替代默认的TkAgg后端。')

add_para('未来提升规划：(1) 学习XGBoost和LightGBM等梯度提升树模型，它们在实际竞赛和工业应用中通常优于随机森林；(2) 了解深度学习（如MLP/TabNet）在表格数据回归预测中的应用场景和优劣势；(3) 加强统计学理论基础（假设检验、贝叶斯推断等），提升数据分析的严谨性；(4) 学习更系统的特征工程方法论（目标编码、特征交叉、自动化特征生成），提升从原始数据中构造高价值特征的能力。')

# ── 保存 ──
doc.save(OUTPUT)
print(f"报告已生成: {OUTPUT}")
